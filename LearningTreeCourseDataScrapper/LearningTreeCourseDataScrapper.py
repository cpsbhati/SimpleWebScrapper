from bs4 import BeautifulSoup
import requests
import re
from docx import Document
import tkinter as tk
from tkinter import messagebox

def scrape_learning_tree_course(url):
    # Check if the URL is from Learning Tree
    if not re.match(r'https?://(www\.)?learningtree\.com/courses/.+', url):
        messagebox.showerror("Error", "Please enter a valid Learning Tree course page URL.")
        return

    try:
        # Send a GET request to the URL
        response = requests.get(url)
        response.raise_for_status()  # Raise an exception for HTTP errors

        # Parse the HTML content
        soup = BeautifulSoup(response.content, "html.parser")

        # Find the elements with specified class names
        course_name = soup.find(class_="page-hero__title")
        course_hero_details = soup.find_all(class_="page-hero-list")
        course_dates = soup.find_all(class_="courses-date__item-title")
        course_details = soup.find_all(class_="course-detail-info__content-outline col-8")

        # Create a new Word document
        doc = Document()

        # Add course name
        if course_name:
            doc.add_heading(f"-------- {course_name.text.strip()} ---------", level=1)

        # Add course hero details
        doc.add_heading("-------- Course Summary Start ---------", level=2)
        for detail in course_hero_details:
            doc.add_paragraph(detail.text.strip())
        doc.add_heading("-------- Course Summary End ---------", level=2)

        # Add course dates
        doc.add_heading("-------- Course Dates Start ---------", level=2)
        for date in course_dates:
            date_text = date.text.strip().replace('\n', " # ").replace("Guaranteed to Run - you can rest assured that the class will not be cancelled. #  # ", "")
            doc.add_paragraph(date_text)
        doc.add_heading("-------- Course Dates End ---------", level=2)

        # Add course details
        doc.add_heading("-------- Course Details ---------", level=2)
        for detail in course_details:
            doc.add_paragraph(detail.text.strip())
        doc.add_heading("-------- Course Details End ---------", level=2)

        # Save the document
        doc.save('scraped_data.docx')
        messagebox.showinfo("Success", "Scraped data written to the scraped_data.docx file")

    except requests.exceptions.RequestException as e:
        messagebox.showerror("Error", f"Failed to get the data from {url}: {e}")

# Create a tkinter GUI window
def get_url_and_scrape():
    url = url_entry.get()
    scrape_learning_tree_course(url)

root = tk.Tk()
root.title("Learning Tree Course Scraper")

# URL Entry
url_label = tk.Label(root, text="Enter Learning Tree course URL:")
url_label.pack(pady=10)
url_entry = tk.Entry(root, width=50)
url_entry.pack(pady=10)

# Scrape Button
scrape_button = tk.Button(root, text="Scrape Course Data", command=get_url_and_scrape)
scrape_button.pack(pady=10)

# Run the GUI
root.mainloop()
