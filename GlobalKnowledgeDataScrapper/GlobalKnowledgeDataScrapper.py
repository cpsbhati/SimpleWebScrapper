from bs4 import BeautifulSoup
import requests
import re
from docx import Document
import tkinter as tk
from tkinter import messagebox

def scrap_data_from_web(file_name, url):
    try:
        # Send a GET request to the URL
        response = requests.get(url)
        response.raise_for_status()  # Raise an HTTPError for bad responses

        # Parse the content with BeautifulSoup
        soup = BeautifulSoup(response.content, 'html.parser')

        # Collect course details
        course_name = soup.find('h1').text.strip()
        course_desc = soup.find(class_="editor").find('p').text.strip()

        # Find the <script> tag(s)
        script_tags = soup.find_all('script')

        # Define a regular expression to match date ranges
        date_pattern = re.compile(
            r'Virtual Classroom Live \| [A-Za-z]{3} \d{2} - (?:[A-Za-z]{3} )?\d{2}, \d{4} \| \d{1,2}:\d{2} [APM]{2} - \d{1,2}:\d{2} [APM]{2} [A-Z]{3} \| ONLINE'
        )

        # Initialize an empty list to collect all event dates
        all_events = []

        # Extract events from each <script> tag
        for script in script_tags:
            script_content = script.string
            if script_content:
                events = date_pattern.findall(script_content)
                if events:
                    all_events.extend(events)

        # Create a Word document
        doc = Document()
        file_save_name = course_name + ' ' + file_name

        # Add course name as title
        doc.add_heading(course_name, level=0)
    
        # Add course description
        doc.add_heading("Course description", level=1)
        doc.add_paragraph(course_desc + '\n')

        # Add dates as bullet points
        for date in all_events:
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            run = p.add_run(date.replace("Virtual Classroom Live | ", "").replace(" | ONLINE", ""))
   
        # Save the Word file
        doc.save(file_save_name)

        messagebox.showinfo("Scraping Complete", f"Data scraped from {url} and saved to {file_save_name}")
    except requests.exceptions.RequestException as e:
        messagebox.showerror("Error", f"Error fetching data from {url}: {e}")
    except AttributeError as e:
        messagebox.showerror("Error", f"Error parsing HTML content: {e}")

def scrape_and_save():
    url = url_entry.get().strip()
    if not re.match(r'https?://(www\.)?globalknowledge\.com/us-en/course/.+', url):
        messagebox.showerror("Error","Please enter a valid Global Knowledge course page URL.")
        return
    if not url:
        messagebox.showerror("Error", "Please enter a valid URL.")
        return
    
    file_name = "course_details.docx"
    scrap_data_from_web(file_name, url)

# Create a tkinter window
root = tk.Tk()
root.title("GK Course Page Scraping Tool")

# Create URL entry and label
url_label = tk.Label(root, text="Enter URL:")
url_label.pack(pady=10)
url_entry = tk.Entry(root, width=50)
url_entry.pack()

# Create scrape button
scrape_button = tk.Button(root, text="Scrape and Save", command=scrape_and_save)
scrape_button.pack(pady=20)

# Run the tkinter main loop
root.mainloop()
